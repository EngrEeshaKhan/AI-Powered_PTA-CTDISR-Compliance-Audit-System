from pathlib import Path

from docx import Document


def parse_docx(file_path: str | Path) -> dict:
    """
    Extract text from a DOCX document.

    Returns:
        Dictionary containing extracted text and basic statistics.
    """

    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"DOCX file not found: {path}")

    document = Document(path)

    paragraphs = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()

        if text:
            paragraphs.append(text)

    # Also extract text from tables.
    table_text = []

    for table in document.tables:
        for row in table.rows:
            cells = []

            for cell in row.cells:
                cell_text = cell.text.strip()

                if cell_text:
                    cells.append(cell_text)

            if cells:
                table_text.append(" | ".join(cells))

    all_text = paragraphs + table_text

    text = "\n".join(all_text).strip()

    return {
        "text": text,
        "page_count": None,
        "word_count": len(text.split()),
        "character_count": len(text),
    }